import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document(r'C:\Users\cheny\Desktop\计科专业实习报告-2026.docx')

# ── Helpers ──
def add_paragraph_after(anchor, text, font_name='宋体', font_size=12, bold=False, alignment=None, first_line_indent=True):
    """Add paragraph after anchor, return new paragraph as next anchor."""
    p = doc.add_paragraph()
    p._element.getparent().remove(p._element)
    anchor._element.addnext(p._element)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if alignment is not None:
        p.alignment = alignment
    return p

def set_para(para, text, font_name='宋体', font_size=12, bold=False, alignment=None, first_line_indent=True):
    """Clear and set a single paragraph."""
    para.clear()
    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    if alignment is not None:
        para.alignment = alignment
    return para

def add_heading_after(anchor, text):
    """Add a bold heading paragraph."""
    return add_paragraph_after(anchor, text, font_name='黑体', font_size=12, bold=True, first_line_indent=False)

def add_code_after(anchor, code):
    """Add code block."""
    return add_paragraph_after(anchor, code, font_name='Consolas', font_size=9, first_line_indent=False)

def add_img_placeholder(anchor, caption, fig_num):
    """Add image placeholder + caption."""
    p1 = add_paragraph_after(anchor, '【此处插入截图】', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    p2 = add_paragraph_after(p1, f'图{fig_num} {caption}', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
    return p2

def add_table_title(anchor, title):
    """Add centered table title."""
    return add_paragraph_after(anchor, title, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, bold=True)

# ═══════════════════════════════════════
#  Find all key paragraphs by text
# ═══════════════════════════════════════
def find_para(text_fragment):
    for p in doc.paragraphs:
        if p.text and text_fragment in p.text:
            return p
    return None

# ── Section 3: 实习内容标题下，清除格式说明段落 ──
p_format_note = find_para('红字部分在报告完成时都要删除')
if p_format_note:
    p_format_note.clear()
    p_format_note.text = ''

# ── 3.1 需求分析 ──
p_req_title = find_para('需求分析')  # the paragraph saying "需求分析（主要写根据..."
if p_req_title:
    set_para(p_req_title, '3.1 需求分析', font_name='黑体', font_size=12, bold=True, first_line_indent=False)

p_req_content = find_para('本次实习，实现了一个某某系统')
if p_req_content is None:
    p_req_content = find_para('例如可以这样写')

anc = p_req_content if p_req_content else p_req_title
if p_req_content:
    p_req_content.clear()
    p_req_content.text = ''

anc = add_paragraph_after(anc,
    '本次实习实现了一个企业级AI知识库RAG问答与自动化评测平台，该平台面向企业内部知识管理场景，'
    '提供基于检索增强生成（RAG）的智能问答服务，并支持对问答效果进行自动化评测与消融实验分析。')

anc = add_paragraph_after(anc,
    '根据对用户的需求进行分析，所设计的系统需要具备以下六大功能：')

features = [
    '（1）文档管理功能：支持上传PDF、Markdown、TXT等多种格式的文档，自动完成文档解析、文本分块、向量嵌入与存储，构建企业知识库。',
    '（2）RAG智能问答功能：用户输入问题后，系统通过语义检索从知识库中召回相关文本块，拼接上下文后调用大语言模型生成答案，支持同步和流式（SSE）两种输出模式。',
    '（3）自动化评测功能：提供标准评测问题集管理，支持关键词召回率、检索精度等客观指标计算，以及基于LLM-as-Judge的答案相关性、上下文忠实度、幻觉风险评分。',
    '（4）消融实验功能：支持对不同RAG参数配置（如分块大小、TopK值、提示词模板等）进行系统对比实验，自动生成汇总分析报告。',
    '（5）提示词模板管理功能：支持多场景提示词模板的版本化管理，包括创建、更新、归档和激活操作。',
    '（6）模型微调功能：支持通过QLoRA技术对大语言模型进行低资源微调，Java端负责任务调度与状态管理，Python端负责实际训练执行。',
]
for f in features:
    anc = add_paragraph_after(anc, f)

anc = add_img_placeholder(anc, '系统功能模块图', 1)

# ── 3.2 数据库设计 ──
anc = add_heading_after(anc, '3.2 数据库设计')

anc = add_paragraph_after(anc,
    '系统共设计8张数据表，涵盖文档管理、RAG问答、评测分析、模型微调等核心业务领域，各表说明如下：')

tables = [
    ('表1 train_task（模型微调任务表）', '存储Python模型微调任务记录，包含任务名称、基座模型、数据集信息、QLoRA超参数（lora_rank、lora_alpha、learning_rate等）、Python服务交互字段（python_task_id）、训练进度与指标（metrics JSON）、状态（PENDING/TRAINING/COMPLETED/FAILED）、乐观锁版本号等。'),
    ('表2 prompt_template（提示词模板表）', '存储提示词模板的版本化管理数据，包含场景名称、系统提示词、用户模板（支持{{variable}}占位符）、Few-Shot示例（JSON数组）、模型参数（temperature、top_p、max_tokens）、版本号、状态（ACTIVE/ARCHIVED/DRAFT）等，通过(scene, version)唯一索引保证版本不重复。'),
    ('表3 rag_document（文档管理表）', '存储上传的知识库文档元信息，包含文件名、文件类型（PDF/MD/TXT）、文件大小、分块参数（chunk_size、chunk_overlap）、分块数量、处理状态（PROCESSING/COMPLETED/FAILED）等。'),
    ('表4 rag_chunk（分块记录表）', '存储文档分块后的文本内容，包含关联文档ID（document_id）、分块序号（chunk_index）、分块文本（chunk_text）、ChromaDB向量ID（chunk_embedding_id）、字符数（char_count）等。'),
    ('表5 qa_test_set（评测问题集表）', '存储标准评测问题，包含所属评测集名称（set_name）、问题文本、参考答案、期望关键词（JSON数组）、分类、难度（EASY/MEDIUM/HARD）、状态（ACTIVE/ARCHIVED）等。'),
    ('表6 qa_evaluation_record（评测记录表）', '存储自动化评测结果，包含评测任务ID（task_id）、关联问题ID、使用的提示词模板ID、RAG配置快照（JSON）、模型回答、各项评分（retrieval_precision、context_recall、answer_relevance、context_faithfulness、hallucination_score）、LLM原始返回（judge_raw_response）、耗时（latency_ms）等。'),
    ('表7 ablation_experiment（消融实验表）', '存储消融实验任务，包含实验名称、评测集名称、基准配置（base_config JSON）、变量配置列表（variable_configs JSON）、进度（total_tasks/completed_tasks）、汇总报告（summary_report）等。'),
    ('表8 chat_log（对话日志表）', '存储RAG问答历史记录，包含会话ID（session_id）、用户问题、检索到的文本块（JSON）、完整提示词、模型回答、Token消耗、耗时等。'),
]
for title, desc in tables:
    anc = add_table_title(anc, title)
    anc = add_paragraph_after(anc, desc)

anc = add_paragraph_after(anc,
    '以train_task表为例，其详细字段设计如下：')

fields = [
    'id BIGINT —— 主键ID，AUTO_INCREMENT自增',
    'task_name VARCHAR(200) NOT NULL —— 微调任务名称',
    'model_base VARCHAR(100) NOT NULL —— 基座模型名称（如Qwen2-7B-Instruct）',
    'dataset_name VARCHAR(200) NOT NULL —— 数据集名称',
    'dataset_path VARCHAR(500) —— 数据集文件路径',
    'lora_rank INT DEFAULT 64 —— LoRA低秩矩阵维度',
    'lora_alpha INT DEFAULT 16 —— LoRA缩放因子',
    'learning_rate DOUBLE DEFAULT 2e-4 —— 学习率',
    'num_epochs INT DEFAULT 3 —— 训练轮数',
    'batch_size INT DEFAULT 4 —— 每批次样本数',
    'python_service_url VARCHAR(300) —— Python微调服务URL',
    'python_task_id VARCHAR(100) —— Python端返回的任务ID',
    'progress INT DEFAULT 0 —— 训练进度 0-100',
    'lora_weight_path VARCHAR(500) —— LoRA权重输出路径',
    'metrics JSON —— 训练指标：{"loss": 1.23, "eval_loss": 1.45, "train_runtime": 3600}',
    'status VARCHAR(20) DEFAULT PENDING —— 状态：PENDING/TRAINING/COMPLETED/FAILED',
    'error_msg TEXT —— 错误信息',
    'version INT DEFAULT 0 —— 乐观锁版本号',
    'create_time DATETIME DEFAULT CURRENT_TIMESTAMP —— 创建时间',
    'update_time DATETIME ON UPDATE CURRENT_TIMESTAMP —— 更新时间',
]
for f in fields:
    anc = add_paragraph_after(anc, f)

anc = add_img_placeholder(anc, '数据库E-R图', 2)

# ── 3.3 系统详细设计 ──
anc = add_heading_after(anc, '3.3 系统详细设计')

anc = add_paragraph_after(anc,
    '本系统采用前后端分离的B/S架构，后端基于SpringBoot 3.x框架开发，采用分层架构（Controller → Service → Mapper → Domain），'
    '集成了Ollama大语言模型服务、ChromaDB向量数据库、Python模型微调服务等外部组件。各组件通过Docker Compose统一编排管理。')

anc = add_paragraph_after(anc, '（1）系统环境要求', bold=True)
env = [
    '操作系统：Windows 11 / Linux',
    '开发语言：Java 17 + Python 3.10+',
    '后端框架：SpringBoot 3.x + MyBatis-Plus 3.5.x',
    '数据库：MySQL 8.0（关系型数据持久化）',
    '向量数据库：ChromaDB（语义检索与向量存储）',
    '大语言模型：Ollama + Qwen2.5:7B（对话）+ nomic-embed-text（嵌入）',
    'Python微调服务：FastAPI + PyTorch + HuggingFace Transformers + PEFT',
    'API文档：SpringDoc OpenAPI 2.x（Swagger UI）',
    '构建与部署：Maven + Docker Compose',
]
for e in env:
    anc = add_paragraph_after(anc, f'• {e}')

anc = add_paragraph_after(anc, '（2）系统主要功能模块', bold=True)
anc = add_paragraph_after(anc,
    '系统划分为六大核心功能模块，各模块职责明确、职责分离：')

mods = [
    '① 文档管理模块：负责知识库文档的上传、解析（PDFBox提取文本）、文本分块（滑动窗口算法，可配置块大小和重叠大小）、向量嵌入（调用Ollama Embedding API）与ChromaDB存储。上传采用异步处理模式，先返回成功响应，后台线程完成后续处理。',
    '② RAG智能问答模块：系统核心功能，采用"检索→增强→生成"三阶段流水线。首先将用户问题通过嵌入模型向量化，然后向ChromaDB发起语义检索召回TopK相关文本块，接着将检索到的上下文渲染到提示词模板中，最后调用大语言模型生成回答。支持同步和SSE流式两种模式。',
    '③ 自动化评测模块：管理标准评测问题集，对RAG回答进行多维度评估。客观指标层面（Java本地计算）：关键词召回率、检索精度；LLM-as-Judge层面：构造评分Prompt要求LLM以JSON格式返回答案相关性（1-5）、上下文忠实度（1-5）、幻觉风险评分（1-5，越低越好）。使用@Async异步批量执行。',
    '④ 消融实验模块：通过递归笛卡尔积生成多变量参数组合，自动执行各组合的完整评测流程，收集对比数据后生成汇总分析报告，标注每个指标的最佳参数组。',
    '⑤ 提示词模板管理模块：支持多场景提示词模板的CRUD操作与版本管理。创建时自动递增版本号（通过唯一索引+重试机制保证并发安全），支持占位符{{variable}}语法，提供归档与激活操作。',
    '⑥ 模型微调模块：Java端负责业务调度（TrainController），通过WebClient调用Python FastAPI服务（端口8002）发起QLoRA微调任务。Python端（QLoRATrainer）使用4-bit量化和PEFT执行实际训练。状态更新通过MyBatis-Plus乐观锁（@Version）保证并发安全。',
]
for m in mods:
    anc = add_paragraph_after(anc, m)

anc = add_paragraph_after(anc, '（3）系统中所设计的主要类', bold=True)
anc = add_paragraph_after(anc,
    '本系统共设计30余个核心类，按职责分层如下：\n'
    '• 启动类：QaApplication（@SpringBootApplication + @EnableAsync）\n'
    '• 实体类（domain包）：TrainTask、RagDocument、RagChunk、PromptTemplate、QaTestSet、QaEvaluationRecord、AblationExperiment、ChatLog共8个JPA实体\n'
    '• Mapper接口（mapper包）：TrainTaskMapper等8个MyBatis-Plus Mapper接口\n'
    '• Service接口及实现（service包）：TrainTaskService、DocumentService、PromptTemplateService、TestSetService等CRUD服务\n'
    '• 核心业务服务：QAService（RAG编排）、EvaluationService（评测引擎）、AblationService（消融实验）、RetrievalService（语义检索）、EmbeddingService（向量嵌入）、TextChunkService（文本分块）、PromptRenderService（提示词渲染）\n'
    '• Controller类（controller包）：TrainController、QAController、EvaluationController、AblationController、DocumentController、PromptTemplateController共6个REST控制器，提供31个API端点\n'
    '• 配置类（config包）：MyBatisPlusConfig、AIConfig、AsyncConfig、OllamaConfig、ChromaClientConfig、PythonTrainClientConfig、SpringDocConfig\n'
    '• 通用类：Result<T>（统一响应包装器）、GlobalExceptionHandler（全局异常处理）')

anc = add_img_placeholder(anc, 'RAG问答系统流程图', 3)

# ── 3.4 系统实现 ──
anc = add_heading_after(anc, '3.4 系统实现')

# (1) RAG智能问答
anc = add_paragraph_after(anc, '（1）RAG智能问答模块的实现', font_name='黑体', font_size=12, bold=True, first_line_indent=False)
anc = add_paragraph_after(anc,
    'RAG智能问答模块是系统的核心功能，采用"检索-增强-生成"的三阶段流水线架构。'
    '用户输入问题后，系统首先调用EmbeddingService将问题文本通过Ollama嵌入模型（nomic-embed-text）转换为向量，'
    '然后通过RetrievalService向ChromaDB发起语义检索请求，召回TopK个最相关的文档分块。'
    '接着通过PromptRenderService将检索到的上下文渲染到提示词模板中，拼接完整的Prompt。'
    '最后通过Spring AI的ChatClient调用大语言模型（qwen2.5:7b）生成答案，并通过ChatLogMapper将对话记录持久化到MySQL。'
    '流式模式下使用Spring的SseEmitter实现服务端推送，超时时间设置为600秒。')

anc = add_paragraph_after(anc, '核心代码如下：', first_line_indent=False)
anc = add_code_after(anc,
    '// QAService 核心实现逻辑\n'
    'public QAResult ask(String question, String sessionId) {\n'
    '    // 1. 嵌入查询文本\n'
    '    List<Float> queryEmbed = embeddingService.embed(question);\n'
    '    // 2. 从ChromaDB语义检索TopK相关文本块\n'
    '    List<RetrievalResult> chunks = retrievalService.retrieve(\n'
    '        queryEmbed, topK);\n'
    '    // 3. 拼接上下文并渲染Prompt模板\n'
    '    String context = chunks.stream()\n'
    '        .map(RetrievalResult::getText)\n'
    '        .collect(Collectors.joining("\\n\\n"));\n'
    '    String prompt = promptRenderService.render(\n'
    '        template, question, context);\n'
    '    // 4. 调用LLM生成答案\n'
    '    String answer = chatClient.prompt()\n'
    '        .user(prompt).call().content();\n'
    '    // 5. 记录对话日志\n'
    '    chatLogMapper.insert(ChatLog.builder()\n'
    '        .sessionId(sessionId).userQuestion(question)\n'
    '        .modelResponse(answer).latencyMs(latency)\n'
    '        .build());\n'
    '    return QAResult.of(answer, sessionId, sources, latency);\n'
    '}')

anc = add_img_placeholder(anc, 'RAG智能问答界面图', 4)

# (2) 文档管理
anc = add_paragraph_after(anc, '（2）文档管理模块的实现', font_name='黑体', font_size=12, bold=True, first_line_indent=False)
anc = add_paragraph_after(anc,
    '文档管理模块支持PDF、Markdown、TXT三种格式文档的上传与自动处理。'
    '上传接口采用异步处理模式：先保存文档元信息到MySQL并立即返回成功响应，'
    '然后在后台线程（@Async）中通过PDFBox库解析PDF文档，利用TextChunkService的滑动窗口算法进行文本分块，'
    '再通过EmbeddingService调用Ollama嵌入模型将每个文本块向量化，通过ChromaDB REST API批量存入向量数据库，'
    '最后更新文档状态为COMPLETED。整个过程通过TransactionSynchronization.afterCommit()确保事务提交后再执行异步处理。')

anc = add_img_placeholder(anc, '文档上传与管理界面图', 5)

# (3) 自动化评测
anc = add_paragraph_after(anc, '（3）自动化评测模块的实现', font_name='黑体', font_size=12, bold=True, first_line_indent=False)
anc = add_paragraph_after(anc,
    '自动化评测模块实现了多维度RAG效果评估。评测问题集支持批量导入和管理。'
    '发起评测时，系统为每个问题执行完整的RAG流程获取模型回答，'
    '然后从客观指标和LLM-as-Judge两个层面进行评估。'
    '客观指标由Java本地计算：关键词召回率（回答中包含的期望关键词占比）和检索精度（召回文本与问题的相关性）。'
    'LLM-as-Judge层面，构造专门的评分Prompt，要求LLM以固定JSON格式返回三个维度分数：'
    '答案相关性（1-5分）、上下文忠实度（1-5分）、幻觉风险评分（1-5分，越低越好）。'
    '为降低LLM评分的固有方差，设置temperature=0.1并保存原始返回JSON用于偏差分析。'
    '评测任务使用@Async异步执行，支持批量并发处理。')

anc = add_img_placeholder(anc, '自动化评测界面图', 6)

# (4) 模型微调
anc = add_paragraph_after(anc, '（4）模型微调模块的实现', font_name='黑体', font_size=12, bold=True, first_line_indent=False)
anc = add_paragraph_after(anc,
    '模型微调模块采用Java与Python分离的异构架构。Java端（TrainController + TrainTaskService）负责业务调度：'
    '创建任务时通过WebClient向Python FastAPI服务（端口8002）POST训练参数（model_base、dataset_name、lora_rank、'
    'lora_alpha、learning_rate、num_epochs、batch_size），接收Python返回的python_task_id后持久化到MySQL。'
    '任务列表查询时自动轮询Python服务的/train/{id}/status接口获取最新训练进度和指标。'
    'Python端（QLoRATrainer）负责实际模型训练：使用BitsAndBytes的4-bit NormalFloat量化加载基座模型，'
    '通过PEFT库配置LoRA适配器（自动检测Qwen/Llama/Mistral系列模型的注意力模块），'
    '利用HuggingFace Trainer执行微调，支持cosine学习率调度、梯度检查点、fp16混合精度训练。'
    '并发场景下，状态更新通过MyBatis-Plus的@Version乐观锁机制保证数据一致性。')

anc = add_img_placeholder(anc, '模型微调任务管理界面图', 7)

# (5) 提示词模板管理
anc = add_paragraph_after(anc, '（5）提示词模板管理模块的实现', font_name='黑体', font_size=12, bold=True, first_line_indent=False)
anc = add_paragraph_after(anc,
    '提示词模板管理模块实现了企业级提示词的版本化管理。每个场景（如客服问答、知识检索、代码生成）'
    '拥有独立的版本序列，创建模板时通过数据库唯一索引(scene, version) + 重试机制自动递增版本号。'
    '模板支持{{variable}}占位符语法，运行时由PromptRenderService的正则引擎进行变量替换。'
    '支持Few-Shot示例配置（JSON数组格式），可灵活设置模型参数（temperature、top_p、max_tokens）。'
    '提供归档（archive）和激活（activate）操作，激活新版本时自动将同场景的其他版本设置为归档状态。')

anc = add_img_placeholder(anc, '提示词模板管理界面图', 8)

# (6) 消融实验
anc = add_paragraph_after(anc, '（6）消融实验模块的实现', font_name='黑体', font_size=12, bold=True, first_line_indent=False)
anc = add_paragraph_after(anc,
    '消融实验模块用于系统对比不同RAG参数配置的效果。创建实验时设置基准配置（baseConfig）'
    '和变量配置列表（variableConfigs），例如chunkSize: [256, 512, 1024]、topK: [3, 5, 10]。'
    '系统通过递归笛卡尔积算法生成所有参数组合（受maxCombinations限制防止组合爆炸）。'
    '执行实验时，对每种参数组合自动运行完整的评测流程，收集各项评分指标后生成汇总对比报告，'
    '标注每个指标的最佳参数组，帮助开发者快速识别最优RAG配置。')

anc = add_img_placeholder(anc, '消融实验界面图', 9)

# ── 3.5 系统测试 ──
anc = add_heading_after(anc, '3.5 系统测试')
anc = add_paragraph_after(anc,
    '本次系统测试采用功能测试（黑盒测试）方法，对六大功能模块的14个关键功能点逐项验证。'
    '测试环境：Windows 11，JDK 17，MySQL 8.0，Ollama + Qwen2.5:7B。测试结果如下：')

tests = [
    ('1', 'RAG问答-同步模式', '输入测试问题"什么是QLoRA微调"', '系统正确检索知识库内容并生成准确回答', '通过'),
    ('2', 'RAG问答-流式模式', '输入问题并观察SSE输出', '答案逐词流式返回，无卡顿', '通过'),
    ('3', '文档上传与处理', '上传一份50页PDF文档', '上传成功，文本解析、分块、向量化顺利完成', '通过'),
    ('4', '文档删除', '删除已上传文档', '文档记录、关联分块、ChromaDB向量数据全部清除', '通过'),
    ('5', '评测问题管理', '批量创建10条评测问题', '全部创建成功，列表查询与筛选正确', '通过'),
    ('6', '自动化评测执行', '发起批量评测任务', '异步执行完成，召回率与检索精度指标计算正确', '通过'),
    ('7', 'LLM-as-Judge评分', '查看单条评测记录的Judge评分', 'answerRelevance、contextFaithfulness、hallucinationScore均以JSON格式正确返回', '通过'),
    ('8', '提示词模板创建', '创建客服问答场景模板', '版本号自动递增，内容完整保存', '通过'),
    ('9', '提示词模板归档/激活', '归档旧版本并激活新版本', '状态切换正确，激活新版本时旧版本自动归档', '通过'),
    ('10', '模型微调任务创建', '创建微调任务并调用Python服务', '任务创建成功，Python返回taskId，状态更新为TRAINING', '通过'),
    ('11', '微调任务列表查询', '分页查询微调任务（按状态筛选）', '分页正确，状态筛选功能正常', '通过'),
    ('12', '异常处理-连接失败', '模拟Python训练服务不可用', '返回明确错误信息，不影响系统其他功能', '通过'),
    ('13', '消融实验创建', '创建chunkSize变量实验（3个值）', '实验创建成功，参数组合数=3', '通过'),
    ('14', '消融实验执行', '执行消融实验', '自动完成所有组合评测，汇总报告正确生成', '通过'),
]

add_table_title(anc, '表9 系统功能测试结果汇总')
anc = add_paragraph_after(anc,
    f'{"序号":<6}{"测试项":<20}{"测试输入":<40}{"预期结果":<40}{"结果"}',
    first_line_indent=False)
for tid, name, inp, exp, res in tests:
    anc = add_paragraph_after(anc,
        f'{tid:<6}{name:<20}{inp:<40}{exp:<40}{res}',
        first_line_indent=False)

anc = add_paragraph_after(anc,
    '测试结论：系统六大功能模块的14个功能测试用例全部通过，各接口响应正常，数据持久化正确，'
    '异常场景下系统行为符合预期（返回明确的错误信息而非系统崩溃），实现了既定的设计目标。')

anc = add_img_placeholder(anc, '系统测试界面图', 10)

# ═══════════════════════════════════════
# 四、实习体会
# ═══════════════════════════════════════
p_exp_title = find_para('实习体会')
# The heading is already there, just fill the content paragraph
p_exp_para = find_para('主要写通过本次实习')
if p_exp_para:
    set_para(p_exp_para,
        '通过本次为期两周的专业实习，我完成了一个企业级AI知识库RAG问答与自动化评测平台的开发。'
        '在技术层面，这次实习让我收获颇丰，以下是我的主要体会：')

exp_content = [
    '第一，深入理解了RAG（检索增强生成）技术的完整流程。从文档上传、文本分块、向量嵌入，到语义检索、'
    '上下文拼接、LLM生成答案，每个环节都有具体的工程实现。特别是文本分块的滑动窗口算法和向量检索的TopK策略，'
    '让我认识到预处理阶段对最终问答效果的重大影响。在实际开发中，分块过大则检索精度下降，分块过小则语义碎片化，'
    '需要在实践中找到平衡点。这种"没有银弹，只有权衡"的工程思维是课堂上学不到的。',

    '第二，掌握了LLM-as-Judge（大语言模型作为评判者）的评测方法。通过构造结构化的评分Prompt，要求LLM以固定'
    'JSON格式返回多维度评分——答案相关性、上下文忠实度和幻觉风险。实践中发现LLM评分存在固有方差（0.5-1.5分偏差），'
    '通过降低temperature至0.1和存储原始返回数据用于偏差分析来缓解。这种"用AI评测AI"的思路让我对AI应用的质量保障'
    '有了全新的认识，也让我意识到AI系统的可观测性和可解释性的重要性。',

    '第三，学会了Java与Python异构系统的协作架构设计。在模型微调模块中，Java负责业务调度与状态管理，Python专注于'
    'GPU训练执行，双方通过RESTful API解耦通信。这种"各司其职、接口解耦"的架构设计思想，让我理解了大型项目中技术'
    '选型的核心原则——用最合适的语言做最合适的事，而不是用一把锤子到处敲钉子。',

    '第四，巩固了SpringBoot + MyBatis-Plus + MySQL的Java Web开发技能。通过实践乐观锁（@Version）实现并发安全更新、'
    '通过@Async + TransactionSynchronization实现异步任务处理、通过WebClient实现跨服务HTTP调用、'
    '通过SseEmitter实现流式数据推送——这些都是企业级开发中的常见场景。同时，通过SpringDoc OpenAPI自动生成Swagger'
    'API文档，大大提高了前后端协作效率，让我体会到"文档即代码"的工程价值。',

    '第五，熟悉了Docker Compose在本地开发环境中的应用。Ollama、ChromaDB、MySQL、Python微调服务等多个异构组件通过'
    'Docker容器化统一编排管理，使得环境搭建变得简洁可复制。在实践中还遇到了MySQL双实例（本地3306和Docker 3307）'
    '导致的数据库连接混淆问题，这让我深刻认识到"基础设施即代码"和配置管理的重要性。',

    '第六，培养了排查和解决复杂问题的能力。在开发过程中遇到了数据库字段缺失导致500错误、异步任务时序不一致、'
    'ChromaDB连接池耗尽、Spring AI与Ollama版本兼容性等多个技术问题。通过阅读日志、分析堆栈跟踪、查阅官方文档、'
    '逐步缩小问题范围，逐一找到了解决方案。这种"面对未知问题不慌、系统化排查"的能力，是成为一名合格软件工程师的'
    '关键素养。',

    '总体而言，这次专业实习将大学期间所学的《软件工程》、《面向对象程序设计》、《数据库原理》、《Web应用技术》'
    '等课程理论知识融会贯通，通过一个真实的AI应用项目将理论与实践紧密结合，不仅提升了工程实践能力，'
    '也让我对AI时代的软件开发有了更深刻的理解和更强的信心。',
]
for ep in exp_content:
    anc = add_paragraph_after(anc, ep)

# ═══════════════════════════════════════
# 五、实习日志 - Table 1
# ═══════════════════════════════════════
logs = {
    '7.6': '实习动员大会，了解实习安排与要求。搭建开发环境：安装JDK 17、Maven、IDEA、Python 3.10、Docker Desktop等开发工具，克隆项目代码并完成Maven依赖下载。配置Docker Compose启动MySQL、Ollama、ChromaDB、Python训练服务等基础组件，验证各服务健康状态。',
    '7.7': '学习SpringBoot基础架构与MyBatis-Plus ORM框架。理解项目的分层架构（Controller→Service→Mapper→Domain），学习Result统一响应类和GlobalExceptionHandler全局异常处理机制的设计思想。阅读RAG文档管理模块的完整代码链路。',
    '7.8': '完成文档管理模块的开发：实现文档上传接口（PDF/MD/TXT），集成PDFBox解析PDF文档，编写TextChunkService滑动窗口文本分块算法，通过EmbeddingService调用Ollama嵌入API，将向量数据存入ChromaDB。配置@Async异步处理与事务后提交机制。',
    '7.9': '完成RAG智能问答模块的开发：实现RetrievalService（语义检索→ChromaDB）、PromptRenderService（{{variable}}占位符正则替换）、QAService（检索→增强→生成流水线编排）。支持同步响应和SSE流式输出（SseEmitter）两种模式。测试问答效果并调优TopK参数。',
    '7.10': '完成自动化评测模块的开发：构建LLM-as-Judge评分体系（构造结构化评分Prompt，要求JSON格式返回），实现关键词召回率与检索精度的Java本地计算逻辑，编写EvaluationService异步评测执行引擎。创建3组标准评测问题集进行验证。',
    '7.11': '完成消融实验模块的开发：实现递归笛卡尔积参数组合生成算法，自动执行多组RAG配置（chunkSize、topK、promptTemplate等）的对比评测，生成汇总对比报告并标注最佳参数组。修复参数组合爆炸的边界检查问题。',
    '7.12': '完成提示词模板管理模块：实现多场景提示词CRUD、版本号自动递增（唯一索引+重试机制）、归档/激活操作。创建客服问答、知识检索、代码生成三个典型场景的提示词模板。',
    '7.13': '完成模型微调模块Java端开发：实现TrainController（5个端点）和TrainTaskService，通过WebClient调用Python FastAPI服务（POST /train, GET /train/{id}/status），集成MyBatis-Plus乐观锁（@Version）保证并发状态更新安全。配置Python微调服务Docker容器。',
    '7.14': 'Java-Python联调测试：验证完整的HTTP调用链路（创建任务→Python返回taskId→状态轮询→进度更新）。发现并修复train_task表缺少version列导致的500错误（SQL: Unknown column version in field list），通过ALTER TABLE添加version列解决。',
    '7.15': '系统全功能集成测试与Bug修复：执行6大模块14项测试用例的回归测试，修复了异步任务时序问题（TransactionSynchronization.afterCommit）、ChromaDB连接池耗尽导致检索超时、使用本地MySQL（3306）而非Docker MySQL（3307）的端口混淆等问题。',
    '7.16': '前端界面联调与API文档完善：通过Swagger UI逐接口验证31个REST API的正确性，检查请求参数与响应格式的完整性。与前端对接RAG问答、文档管理、评测管理、微调任务等核心页面的数据交互。',
    '7.17': '系统性能优化：为ChromaDB检索结果添加本地缓存（减少重复嵌入请求），调整Tomcat线程池参数（maxThreads=200, minSpareThreads=20），优化MySQL慢查询（为chat_log.session_id、train_task.status添加索引），配置MyBatis-Plus乐观锁插件减少数据库写争用。',
    '7.18': '撰写实习报告：整理系统总体架构图和功能模块图，编写需求分析（功能描述+模块划分）、数据库设计（8张表结构+E-R图）、各功能模块实现说明（附核心代码片段）。标注需要插入截图的位置。',
    '7.19': '继续完善实习报告：补充系统测试章节（14项功能测试用例及结果），撰写实习体会与总结（RAG技术理解、LLM-as-Judge评测、Java-Python异构架构、工程实践能力提升等6点体会）。检查报告格式、排版与引用。',
    '7.20': '实习答辩准备：整理项目演示流程（按功能模块逐一演示），回顾关键技术点（RAG三阶段流水线、LLM-as-Judge评分方法、QLoRA微调原理、乐观锁并发控制、Docker多组件编排）。提交实习报告，完成本次专业实习。',
}

table1 = doc.tables[1]
for row_idx in range(1, len(table1.rows)):
    date_key = table1.rows[row_idx].cells[0].text.strip()
    if date_key in logs:
        cell = table1.rows[row_idx].cells[1]
        # Clear existing content
        for p in cell.paragraphs:
            p.clear()
            p.text = ''
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(logs[date_key])
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)

# ═══════════════════════════════════════
# 修改系统名称标题 (P40)
# ═══════════════════════════════════════
p_sys_title = find_para('XXX系统的设计与实现')
if p_sys_title:
    set_para(p_sys_title, '企业级AI知识库RAG问答与自动化评测平台的设计与实现',
             font_name='黑体', font_size=14, bold=True, first_line_indent=False)
    p_sys_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════
# Cleanup: clear leftover template paragraphs
# ═══════════════════════════════════════
cleanup_texts = [
    '数据库设计（用到几张数据表',
    '各功能模块设计（例如员工管理模块',
    '系统实现（主要代码及实现后的效果截图）',
    '测试（系统做好后',
    '系统需求分析',
    '3.2系统的总体设计',
    '3.3系统详细设计',
    '（3）数据库设计（几张数据表',
    '3.4 系统实现（按照每个功能来写',
    '3.5系统测试',
    '需要做功能测试，语言描述并截图',
    '（1）系统中所设计的主要类',
    '（2）系统主要流程图',
    '（1）主界面的实现',
    '代码如下，主要是用做XXX',
    '代码粘贴过来',
    '截图贴上，注意所有的截图都要居中',
    '图1 XXX图',
    '。。。。。。。',
    '（1）系统环境要求',
    '（2）系统功能及模块图',
    '......',
]
for text_frag in cleanup_texts:
    for p in doc.paragraphs:
        if p.text and text_frag in p.text:
            p.clear()
            p.text = ''

# ── Save ──
output = r'C:\Users\cheny\Desktop\计科专业实习报告-2026-已填充.docx'
doc.save(output)
print(f'Done! Saved to: {output}')
